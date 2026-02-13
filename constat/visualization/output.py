# Copyright (c) 2025 Kenneth Stott
#
# This source code is licensed under the Business Source License 1.1
# found in the LICENSE file in the root directory of this source tree.
#
# NOTICE: Use of this software for training artificial intelligence or
# machine learning models is strictly prohibited without explicit written
# permission from the copyright holder.

"""Output helper for saving files and visualizations.

Saves outputs in two ways:
1. To disk for immediate CLI access (file:// URIs users can click)
2. As artifacts in datastore for React UI to display

Usage in generated code:
    # Save a document (markdown, text, etc.)
    viz.save_file('quarterly_report', content, ext='md', title='Q4 Report')

    # Save a folium map
    import folium
    m = folium.Map(location=[50, 10], zoom_start=4)
    viz.save_map('euro_countries', m, title='Countries Using Euro')

    # Save a plotly chart
    import plotly.express as px
    fig = px.bar(df, x='country', y='population')
    viz.save_chart('population_chart', fig, title='Population by Country')

    # Save raw HTML
    viz.save_html('custom_viz', html_string, title='Custom Visualization')
"""

from __future__ import annotations

import json
import os
import subprocess
import sys
from pathlib import Path
from typing import Any, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from constat.storage.datastore import DataStore
    from constat.storage.registry import ConstatRegistry


def is_repl_mode() -> bool:
    """Check if running in REPL mode."""
    return os.environ.get("CONSTAT_REPL_MODE", "").lower() in ("1", "true", "yes")


# Global to collect outputs during REPL execution
_pending_outputs: list[dict] = []


def clear_pending_outputs() -> None:
    """Clear the pending outputs list."""
    global _pending_outputs
    _pending_outputs = []


def get_pending_outputs() -> list[dict]:
    """Get and clear pending outputs for display."""
    global _pending_outputs
    outputs = _pending_outputs[:]
    _pending_outputs = []
    return outputs


def peek_pending_outputs() -> list[dict]:
    """Peek at pending outputs without clearing them."""
    return _pending_outputs[:]


def add_pending_output(file_uri: str, description: str, file_type: str = "") -> None:
    """Add an output to the pending list for later display."""
    _pending_outputs.append({
        "file_uri": file_uri,
        "description": description,
        "type": file_type,
    })


# File extension to artifact type mapping
FILE_EXT_ARTIFACT_TYPES = {
    # Text formats
    "md": "markdown",
    "markdown": "markdown",
    "txt": "text",
    "text": "text",
    "csv": "csv",
    "json": "json",
    "xml": "xml",
    "yaml": "yaml",
    "yml": "yaml",
    "html": "html",
    "htm": "html",
    # Spreadsheets
    "xlsx": "spreadsheet",
    "xls": "spreadsheet",
    "ods": "spreadsheet",
    # Documents
    "pdf": "document",
    "docx": "document",
    "doc": "document",
    "odt": "document",
    "rtf": "document",
    # Presentations
    "pptx": "presentation",
    "ppt": "presentation",
    "odp": "presentation",
    # Images
    "png": "image",
    "jpg": "image",
    "jpeg": "image",
    "gif": "image",
    "svg": "image",
    "bmp": "image",
    "tiff": "image",
    "tif": "image",
    "webp": "image",
    # Data formats
    "parquet": "data",
    "arrow": "data",
    "feather": "data",
    "pickle": "data",
    "pkl": "data",
}


class VisualizationHelper:
    """Helper for saving files and visualizations.

    Provides a simple interface for generated code to save files and
    interactive visualizations. Files are saved to an output directory
    and also registered as artifacts in the datastore for the React UI.

    Attributes:
        output_dir: Directory where files are saved
        datastore: DataStore for artifact registration (optional)
        print_file_refs: Whether to print file:// URIs (True for CLI, False for React UI)
        open_with_system_viewer: Auto-open saved files in system default app if True
        session_id: Session ID for organizing outputs
    """

    def __init__(
        self,
        output_dir: Optional[Path] = None,
        datastore: Optional["DataStore"] = None,
        step_number: int = 0,
        print_file_refs: bool = True,
        session_id: Optional[str] = None,
        user_id: str = "default",
        registry: Optional["ConstatRegistry"] = None,
        open_with_system_viewer: bool = False,
    ):
        """Initialize the output helper.

        Args:
            output_dir: Directory for saving files. Defaults to .constat/<user_id>/artifacts/
            datastore: DataStore for registering artifacts (for React UI)
            step_number: Current step number for artifact metadata
            print_file_refs: If True, print file:// URIs for CLI users.
                           If False, suppress file references (for React UI where
                           artifacts are displayed directly).
            session_id: Session ID for organizing outputs by session
            user_id: User ID for user-scoped storage (default: "default")
            registry: Central registry for artifact tracking (optional)
            open_with_system_viewer: If True, auto-open saved files in the OS default app
        """
        self.session_id = session_id
        self.user_id = user_id
        self.registry = registry
        self.open_with_system_viewer = open_with_system_viewer

        if output_dir is None:
            # Use session directory under user-scoped storage
            # "artifacts" = user-requested outputs (charts, files, visualizations)
            if session_id:
                output_dir = Path(".constat") / user_id / "sessions" / session_id / "artifacts"
            else:
                output_dir = Path(".constat") / user_id / "artifacts"

        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.datastore = datastore
        self.step_number = step_number
        self.print_file_refs = print_file_refs

    def _generate_filename(self, name: str, extension: str) -> Path:
        """Generate a unique filename for the output."""
        # Sanitize name for filesystem
        safe_name = "".join(c if c.isalnum() or c in "-_" else "_" for c in name)
        return self.output_dir / f"{safe_name}.{extension}"

    def _file_uri(self, path: Path) -> str:
        """Convert a path to a file:// URI for clickable terminal links."""
        return path.resolve().as_uri()

    def _open_in_system_viewer(self, filepath: Path) -> None:
        """Open the file in the OS default application (non-blocking).

        Uses platform-appropriate methods:
        - macOS: 'open' command
        - Windows: os.startfile() (avoids shell=True security risk)
        - Linux: xdg-open
        """
        path_str = str(filepath.resolve())
        if sys.platform.startswith("darwin"):
            # macOS
            subprocess.Popen(["open", path_str], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        elif sys.platform.startswith("win"):
            # Windows: use os.startfile() to avoid command injection risk with shell=True
            os.startfile(path_str)  # type: ignore[attr-defined]
        else:
            # Linux / other Unix
            subprocess.Popen(["xdg-open", path_str], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    def _print_ref(self, label: str, filepath: Path, description: str = "") -> None:
        """Print file reference or collect for later display.

        In REPL mode, outputs are collected and displayed in an "Outputs:" section
        at the end of execution rather than printed inline.

        In server/web mode (when datastore is present), prints just the filename
        since artifacts appear in the Results panel. Full paths are hidden for security.
        """
        if not self.print_file_refs and not self.open_with_system_viewer:
            return

        file_uri = self._file_uri(filepath)
        desc = description or label

        # Optionally open with system viewer to bypass IDE interception of file:// links
        if self.open_with_system_viewer:
            try:
                self._open_in_system_viewer(filepath)
            except Exception as e:
                # Non-fatal; continue to print refs if requested
                if self.print_file_refs:
                    print(f"Note: Could not auto-open file in system viewer: {e}")

        if self.print_file_refs:
            # Always collect for REPL display AND print
            # REPL will collect from pending_outputs, CLI will show the print
            add_pending_output(file_uri, desc, filepath.suffix.lstrip("."))
            if not is_repl_mode():
                # In server/web mode, show just filename (artifacts visible in Results panel)
                # In CLI mode without datastore, show full URI for clickable links
                if self.datastore:
                    # Web UI mode - just show filename, hide server paths for security
                    print(f"Saved: {filepath.name}")
                else:
                    # CLI mode - show full file:// URI for clickable terminal links
                    print(f"{label}: {file_uri}")

    def _register_artifact(
        self,
        filepath: Path,
        artifact_type: str,
        description: Optional[str] = None,
    ) -> None:
        """Register an artifact in the central registry.

        Args:
            filepath: Path to the saved file
            artifact_type: Type of artifact (chart, image, csv, etc.)
            description: Human-readable description
        """
        if not self.registry or not self.session_id:
            return

        try:
            size_bytes = filepath.stat().st_size if filepath.exists() else 0
            self.registry.register_artifact(
                user_id=self.user_id,
                session_id=self.session_id,
                name=filepath.name,
                file_path=str(filepath),
                artifact_type=artifact_type,
                size_bytes=size_bytes,
                description=description,
            )
        except Exception:
            pass  # Don't fail the save if registry fails

    def save_file(
        self,
        name: str,
        content: str,
        ext: str = "txt",
        title: Optional[str] = None,
        description: Optional[str] = None,
    ) -> Path:
        """Save any text file to disk and artifact store.

        Use this for documents, reports, data exports, etc.

        Args:
            name: Name for the file (used in filename and artifact)
            content: Text content to save
            ext: File extension (md, txt, csv, json, etc.)
            title: Human-readable title for the artifact
            description: Description of the file

        Returns:
            Path to the saved file
        """
        # Normalize extension
        ext = ext.lstrip(".")

        # Save to file
        filepath = self._generate_filename(name, ext)
        filepath.write_text(content, encoding="utf-8")

        # Register as artifact if datastore available
        if self.datastore:
            try:
                artifact_type = FILE_EXT_ARTIFACT_TYPES.get(ext, "text")
                self.datastore.save_rich_artifact(
                    name=name,
                    artifact_type=artifact_type,
                    content=content,
                    step_number=self.step_number,
                    title=title or name,
                    description=description,
                )
            except Exception as e:
                # Don't fail if artifact save fails - file is already saved
                if self.print_file_refs:
                    print(f"Note: Could not register artifact: {e}")

        # Register in central registry
        artifact_type = FILE_EXT_ARTIFACT_TYPES.get(ext, "text")
        self._register_artifact(filepath, artifact_type, description=title or name)

        self._print_ref(f"File saved ({ext})", filepath, description=title or name)
        return filepath

    def save_binary(
        self,
        name: str,
        content: bytes,
        ext: str = "xlsx",
        title: Optional[str] = None,
        description: Optional[str] = None,
    ) -> Path:
        """Save binary file to disk and artifact store.

        Use this for Excel files, images, PDFs, and other binary formats.

        For XLSX files: generates an HTML table preview for viewing in the UI.
        For other Office docs: saves as HTML with a download link.

        Args:
            name: Name for the file (used in filename and artifact)
            content: Binary content to save
            ext: File extension (xlsx, pdf, png, etc.)
            title: Human-readable title for the artifact
            description: Description of the file

        Returns:
            Path to the saved file
        """
        from io import BytesIO

        # Normalize extension
        ext = ext.lstrip(".")

        # Save to file
        filepath = self._generate_filename(name, ext)
        filepath.write_bytes(content)
        file_uri = filepath.resolve().as_uri()

        # Register as artifact if datastore available
        artifact_type = FILE_EXT_ARTIFACT_TYPES.get(ext, "document")
        if self.datastore:
            try:
                if ext == "xlsx":
                    # XLSX: Generate HTML table preview
                    import pandas as pd
                    try:
                        # Read all sheets
                        xlsx_data = pd.read_excel(BytesIO(content), sheet_name=None)
                        html_parts = [f"<h2>{title or name}</h2>"]
                        for sheet_name, df in xlsx_data.items():
                            if len(xlsx_data) > 1:
                                html_parts.append(f"<h3>Sheet: {sheet_name}</h3>")
                            # Limit preview to first 100 rows
                            preview_df = df.head(100)
                            html_parts.append(preview_df.to_html(index=False, classes="data-table", border=0))
                            if len(df) > 100:
                                html_parts.append(f"<p><em>Showing 100 of {len(df)} rows</em></p>")
                        html_parts.append(f'<p><a href="#" class="constat-download" data-filename="{filepath.name}">Download full file ({ext.upper()})</a></p>')
                        html_content = self._wrap_html_preview("\n".join(html_parts))
                        self.datastore.save_rich_artifact(
                            name=name,
                            artifact_type="html",
                            content=html_content,
                            step_number=self.step_number,
                            title=title or name,
                            description=description,
                            metadata={"file_path": str(filepath), "source_type": "xlsx", "extension": ext},
                        )
                    except Exception as e:
                        # Fall back to download link if preview fails
                        html_content = self._wrap_html_preview(
                            f'<h2>{title or name}</h2>'
                            f'<p>Excel spreadsheet</p>'
                            f'<p><a href="#" class="constat-download" data-filename="{filepath.name}">Download {filepath.name}</a></p>'
                            f'<p><small>Preview unavailable: {e}</small></p>'
                        )
                        self.datastore.save_rich_artifact(
                            name=name,
                            artifact_type="html",
                            content=html_content,
                            step_number=self.step_number,
                            title=title or name,
                            description=description,
                            metadata={"file_path": str(filepath), "source_type": "xlsx", "extension": ext},
                        )
                elif ext in ("docx", "doc"):
                    # DOCX: Convert to HTML with mammoth
                    html_content = self._convert_docx_to_html(content, title or name, file_uri, filepath.name)
                    self.datastore.save_rich_artifact(
                        name=name,
                        artifact_type="html",
                        content=html_content,
                        step_number=self.step_number,
                        title=title or name,
                        description=description,
                        metadata={"file_path": str(filepath), "source_type": "docx", "extension": ext},
                    )
                elif ext == "pdf":
                    # PDF: Embed with object tag for native browser viewing
                    html_content = self._convert_pdf_to_html(content, title or name, file_uri, filepath.name)
                    self.datastore.save_rich_artifact(
                        name=name,
                        artifact_type="html",
                        content=html_content,
                        step_number=self.step_number,
                        title=title or name,
                        description=description,
                        metadata={"file_path": str(filepath), "source_type": "pdf", "extension": ext},
                    )
                else:
                    # Other Office docs (PPTX, etc.): show download link
                    ext_labels = {
                        "pptx": "PowerPoint Presentation",
                        "ppt": "PowerPoint Presentation",
                        "odt": "OpenDocument Text",
                        "ods": "OpenDocument Spreadsheet",
                        "odp": "OpenDocument Presentation",
                    }
                    ext_label = ext_labels.get(ext, f"{ext.upper()} file")
                    html_content = self._wrap_html_preview(
                        f'<h2>{title or name}</h2>'
                        f'<p>{ext_label}</p>'
                        f'<p><a href="#" class="constat-download" data-filename="{filepath.name}">Download {filepath.name}</a></p>'
                    )
                    self.datastore.save_rich_artifact(
                        name=name,
                        artifact_type="html",
                        content=html_content,
                        step_number=self.step_number,
                        title=title or name,
                        description=description,
                        metadata={"file_path": str(filepath), "source_type": ext, "extension": ext},
                    )
            except Exception as e:
                # Don't fail if artifact save fails - file is already saved
                if self.print_file_refs:
                    print(f"Note: Could not register artifact: {e}")

        # Register in central registry
        self._register_artifact(filepath, artifact_type, description=title or name)

        self._print_ref(f"File saved ({ext})", filepath, description=title or name)
        return filepath

    def _convert_docx_to_html(self, content: bytes, title: str, file_uri: str, filename: str) -> str:
        """Convert DOCX to HTML using mammoth library."""
        from io import BytesIO
        try:
            import mammoth
            result = mammoth.convert_to_html(BytesIO(content))
            html_body = result.value
            # Add title and download link
            body_content = (
                f'<h2>{title}</h2>'
                f'<div class="docx-content">{html_body}</div>'
                f'<p><a href="#" class="constat-download" data-filename="{filename}">Download original document</a></p>'
            )
            if result.messages:
                warnings = "<br>".join(str(m) for m in result.messages[:5])
                body_content += f'<p><small>Conversion notes: {warnings}</small></p>'
            return self._wrap_html_preview(body_content)
        except ImportError:
            # mammoth not installed, fall back to python-docx text extraction
            return self._convert_docx_to_html_fallback(content, title, file_uri, filename)
        except Exception as e:
            return self._wrap_html_preview(
                f'<h2>{title}</h2>'
                f'<p>Word Document</p>'
                f'<p><a href="#" class="constat-download" data-filename="{filename}">Download {filename}</a></p>'
                f'<p><small>Preview unavailable: {e}</small></p>'
            )

    def _convert_docx_to_html_fallback(self, content: bytes, title: str, _file_uri: str, filename: str) -> str:
        """Fallback DOCX to HTML using python-docx for text extraction."""
        from io import BytesIO
        try:
            from docx import Document
            doc = Document(BytesIO(content))
            paragraphs = []
            for para in doc.paragraphs[:100]:  # Limit preview
                text = para.text.strip()
                if text:
                    # Simple styling based on paragraph style
                    style = para.style.name.lower() if para.style else ""
                    if "heading 1" in style:
                        paragraphs.append(f"<h3>{text}</h3>")
                    elif "heading" in style:
                        paragraphs.append(f"<h4>{text}</h4>")
                    else:
                        paragraphs.append(f"<p>{text}</p>")
            if len(doc.paragraphs) > 100:
                paragraphs.append(f"<p><em>Showing 100 of {len(doc.paragraphs)} paragraphs</em></p>")
            body_content = (
                f'<h2>{title}</h2>'
                f'<div class="docx-content">{"".join(paragraphs)}</div>'
                f'<p><a href="#" class="constat-download" data-filename="{filename}">Download original document</a></p>'
            )
            return self._wrap_html_preview(body_content)
        except Exception as e:
            return self._wrap_html_preview(
                f'<h2>{title}</h2>'
                f'<p>Word Document</p>'
                f'<p><a href="#" class="constat-download" data-filename="{filename}">Download {filename}</a></p>'
                f'<p><small>Preview unavailable: {e}</small></p>'
            )

    def _convert_pdf_to_html(self, content: bytes, title: str, _file_uri: str, filename: str) -> str:
        """Convert PDF to HTML with embedded viewer using PDF.js."""
        import base64
        # Base64 encode PDF for embedding
        pdf_b64 = base64.b64encode(content).decode("utf-8")
        # Use PDF.js for reliable cross-browser rendering
        return f"""<!DOCTYPE html>
<html>
<head>
<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; padding: 16px; margin: 0; background: #f5f5f5; }}
@media (prefers-color-scheme: dark) {{ body {{ background: #1f2937; color: #e5e7eb; }} }}
h2 {{ margin: 0 0 12px 0; font-size: 1.25rem; }}
.pdf-container {{ background: #fff; border-radius: 8px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }}
canvas {{ display: block; margin: 0 auto; max-width: 100%; }}
.controls {{ padding: 8px 12px; background: #f0f0f0; display: flex; gap: 8px; align-items: center; font-size: 0.875rem; }}
@media (prefers-color-scheme: dark) {{ .controls {{ background: #374151; }} }}
button {{ padding: 4px 12px; border: 1px solid #ccc; border-radius: 4px; background: #fff; cursor: pointer; }}
@media (prefers-color-scheme: dark) {{ button {{ background: #4b5563; border-color: #6b7280; color: #e5e7eb; }} }}
a {{ color: #3b82f6; }}
</style>
</head>
<body>
<h2>{title}</h2>
<div class="pdf-container">
  <div class="controls">
    <button onclick="prevPage()">← Prev</button>
    <span id="page-info">Page 1 of ?</span>
    <button onclick="nextPage()">Next →</button>
    <a href="#" class="constat-download" data-filename="{filename}" style="margin-left: auto;">Download PDF</a>
  </div>
  <canvas id="pdf-canvas"></canvas>
</div>
<script>
pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
const pdfData = atob("{pdf_b64}");
let pdfDoc = null, pageNum = 1, pageRendering = false;
const canvas = document.getElementById('pdf-canvas'), ctx = canvas.getContext('2d');

pdfjsLib.getDocument({{data: pdfData}}).promise.then(pdf => {{
  pdfDoc = pdf;
  document.getElementById('page-info').textContent = 'Page 1 of ' + pdf.numPages;
  renderPage(1);
}});

function renderPage(num) {{
  pageRendering = true;
  pdfDoc.getPage(num).then(page => {{
    const viewport = page.getViewport({{scale: 1.5}});
    canvas.height = viewport.height;
    canvas.width = viewport.width;
    page.render({{canvasContext: ctx, viewport: viewport}}).promise.then(() => {{
      pageRendering = false;
      document.getElementById('page-info').textContent = 'Page ' + num + ' of ' + pdfDoc.numPages;
    }});
  }});
}}

function prevPage() {{ if (pageNum > 1) {{ pageNum--; renderPage(pageNum); }} }}
function nextPage() {{ if (pageNum < pdfDoc.numPages) {{ pageNum++; renderPage(pageNum); }} }}
</script>
</body>
</html>"""

    def _wrap_html_preview(self, body_content: str) -> str:
        """Wrap content in a styled HTML document for artifact preview."""
        return f"""<!DOCTYPE html>
<html>
<head>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; padding: 16px; margin: 0; line-height: 1.6; color: #1f2937; background: #fff; }}
@media (prefers-color-scheme: dark) {{
  body {{ color: #e5e7eb; background: #111827; }}
  a {{ color: #60a5fa; }}
  table {{ border-color: #374151; }}
  th {{ background: #1f2937; }}
}}
h2 {{ margin-top: 0; font-size: 1.25rem; }}
h3 {{ font-size: 1rem; margin-top: 1.5rem; }}
table.data-table {{ border-collapse: collapse; width: 100%; margin: 1em 0; font-size: 0.875rem; }}
table.data-table th, table.data-table td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
table.data-table th {{ background: #f5f5f5; font-weight: 600; }}
table.data-table tr:nth-child(even) {{ background: #fafafa; }}
@media (prefers-color-scheme: dark) {{
  table.data-table th, table.data-table td {{ border-color: #374151; }}
  table.data-table th {{ background: #1f2937; }}
  table.data-table tr:nth-child(even) {{ background: #1f2937; }}
}}
a {{ color: #3b82f6; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}
</style>
</head>
<body>
{body_content}
</body>
</html>"""

    def save_html(
        self,
        name: str,
        html_content: str,
        title: Optional[str] = None,
        description: Optional[str] = None,
    ) -> Path:
        """Save raw HTML content to file and artifact store.

        Args:
            name: Name for the visualization (used in filename and artifact)
            html_content: HTML string to save
            title: Human-readable title for the artifact
            description: Description of the visualization

        Returns:
            Path to the saved HTML file
        """
        # Save to file
        filepath = self._generate_filename(name, "html")
        filepath.write_text(html_content, encoding="utf-8")

        # Register as artifact if datastore available
        if self.datastore:
            try:
                self.datastore.save_html(
                    name=name,
                    html_content=html_content,
                    step_number=self.step_number,
                    title=title or name,
                    description=description,
                )
            except Exception as e:
                # Don't fail if artifact save fails - file is already saved
                if self.print_file_refs:
                    print(f"Note: Could not register artifact: {e}")

        # Register in central registry
        self._register_artifact(filepath, "html", description=title or name)

        self._print_ref("HTML", filepath, description=title or name)
        return filepath

    def save_map(
        self,
        name: str,
        folium_map: Any,
        title: Optional[str] = None,
        description: Optional[str] = None,
    ) -> Path:
        """Save a folium map to file and artifact store.

        Args:
            name: Name for the map (used in filename and artifact)
            folium_map: Folium Map object
            title: Human-readable title for the artifact
            description: Description of the map

        Returns:
            Path to the saved HTML file
        """
        # Get HTML from folium map
        html_content = folium_map._repr_html_()

        # Save to file
        filepath = self._generate_filename(name, "html")
        folium_map.save(str(filepath))

        # Register as artifact if datastore available
        if self.datastore:
            try:
                self.datastore.save_html(
                    name=name,
                    html_content=html_content,
                    step_number=self.step_number,
                    title=title or f"Map: {name}",
                    description=description or "Interactive map visualization",
                )
            except Exception as e:
                if self.print_file_refs:
                    print(f"Note: Could not register artifact: {e}")

        # Register in central registry
        self._register_artifact(filepath, "map", description=title or f"Map: {name}")

        self._print_ref("Map", filepath, description=title or f"Map: {name}")
        return filepath

    def save_chart(
        self,
        name: str,
        figure: Any,
        title: Optional[str] = None,
        _description: Optional[str] = None,
        chart_type: str = "plotly",
    ) -> Path:
        """Save a Plotly or Altair chart to file and artifact store.

        Args:
            name: Name for the chart (used in filename and artifact)
            figure: Plotly Figure or Altair Chart object
            title: Human-readable title for the artifact
            _description: Description of the chart (unused, kept for API compat)
            chart_type: Type of chart ("plotly" or "altair")

        Returns:
            Path to the saved HTML file
        """
        filepath = self._generate_filename(name, "html")

        # Handle different chart libraries
        if chart_type == "plotly" or hasattr(figure, "write_html"):
            # Plotly figure
            figure.write_html(str(filepath), include_plotlyjs=True, full_html=True)
            _html_content = filepath.read_text(encoding="utf-8")

            # Also save chart spec as artifact
            if self.datastore:
                try:
                    spec = json.loads(figure.to_json())
                    self.datastore.save_chart(
                        name=name,
                        spec=spec,
                        step_number=self.step_number,
                        title=title or name,
                        chart_type="plotly",
                    )
                except Exception as e:
                    if self.print_file_refs:
                        print(f"Note: Could not register chart artifact: {e}")

        elif chart_type == "altair" or hasattr(figure, "save"):
            # Altair chart
            figure.save(str(filepath))
            _html_content = filepath.read_text(encoding="utf-8")

            # Also save chart spec as artifact
            if self.datastore:
                try:
                    spec = figure.to_dict()
                    self.datastore.save_chart(
                        name=name,
                        spec=spec,
                        step_number=self.step_number,
                        title=title or name,
                        chart_type="vega-lite",
                    )
                except Exception as e:
                    if self.print_file_refs:
                        print(f"Note: Could not register chart artifact: {e}")
        else:
            raise ValueError(f"Unsupported chart type: {chart_type}")

        # Register in central registry
        self._register_artifact(filepath, "chart", description=title or name)

        self._print_ref("Chart", filepath, description=title or name)
        return filepath

    def save_image(
        self,
        name: str,
        figure: Any,
        fmt: str = "png",
        title: Optional[str] = None,
        _description: Optional[str] = None,
    ) -> Path:
        """Save a matplotlib figure or image to file.

        Args:
            name: Name for the image
            figure: Matplotlib figure or image data
            fmt: Image format (png, svg, jpg)
            title: Human-readable title
            _description: Description of the image (unused, kept for API compat)

        Returns:
            Path to the saved image file
        """
        filepath = self._generate_filename(name, fmt)

        # Handle matplotlib figures
        if hasattr(figure, "savefig"):
            figure.savefig(str(filepath), format=fmt, bbox_inches="tight", dpi=150)
        else:
            # Assume raw bytes
            filepath.write_bytes(figure)

        # Register as artifact if datastore available
        if self.datastore and fmt in ("png", "svg", "jpg", "jpeg"):
            try:
                image_data = filepath.read_bytes()
                from constat.core.models import ArtifactType
                _artifact_type = {
                    "png": ArtifactType.PNG,
                    "svg": ArtifactType.SVG,
                    "jpg": ArtifactType.JPEG,
                    "jpeg": ArtifactType.JPEG,
                }.get(fmt, ArtifactType.PNG)

                self.datastore.save_image(
                    name=name,
                    image_data=image_data,
                    image_format=fmt,
                    step_number=self.step_number,
                    title=title or name,
                )
            except Exception as e:
                if self.print_file_refs:
                    print(f"Note: Could not register image artifact: {e}")

        # Register in central registry
        self._register_artifact(filepath, "image", description=title or name)

        self._print_ref("Image", filepath, description=title or name)
        return filepath


def create_viz_helper(
    datastore: Optional["DataStore"] = None,
    output_dir: Optional[Path] = None,
    step_number: int = 0,
    print_file_refs: bool = True,
    session_id: Optional[str] = None,
    user_id: str = "default",
    registry: Optional["ConstatRegistry"] = None,
    open_with_system_viewer: bool = False,
) -> VisualizationHelper:
    """Create a VisualizationHelper instance.

    Factory function for use in execution globals.

    Args:
        datastore: DataStore for artifact registration
        output_dir: Custom output directory
        step_number: Current step number
        print_file_refs: If True, print file:// URIs (CLI mode).
                        If False, suppress (React UI mode).
        session_id: Session ID for organizing outputs by session
        user_id: User ID for user-scoped storage (default: "default")
        registry: Central registry for artifact tracking
        open_with_system_viewer: If True, open generated files with system viewer

    Returns:
        Configured VisualizationHelper instance
    """
    return VisualizationHelper(
        output_dir=output_dir,
        datastore=datastore,
        step_number=step_number,
        print_file_refs=print_file_refs,
        session_id=session_id,
        user_id=user_id,
        registry=registry,
        open_with_system_viewer=open_with_system_viewer,
    )
